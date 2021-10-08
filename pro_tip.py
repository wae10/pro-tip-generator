import docx
import docx2txt
from docx import Document
from docx.shared import Pt 
from docx.shared import Length
import os

import datetime # for date and time 


def write_to_doc(filepath):
    """Write Pro Tips from class notes to new "Pro Tip" docx

    Args:
        filepath ([string]): [class name (ex: OPIM-258) ]
    """    
    
    class_folder = filepath.replace('/Notes.docx', '')

    # print(class_folder)

    doc = docx.Document('/Users/williameverett/Desktop/Fall-2020/' + filepath)

    results_list = []

    for i in doc.paragraphs:


        if "Pro Tip " in i.text:
            result = i.text

            result = ''.join(map(str, result))

            result = list(result.split(" "))

            result.remove("Pro")
            result.remove("Tip")
            result = ' '.join(result)
            results_list.append(result)
            # print(result)

        if "pro tip " in i.text:
            result = i.text

            result = ''.join(map(str, result))

            result = list(result.split(" "))

            result.remove("pro")
            result.remove("tip")
            result = ' '.join(result)
            results_list.append(result)
            # print(result)

        if "Pro tip " in i.text:
            result = i.text

            result = ''.join(map(str, result))

            result = list(result.split(" "))

            result.remove("Pro")
            result.remove("tip")
            result = ' '.join(result)
            results_list.append(result)
            # print(result)

        if "Pro Tips" in i.text:
            result = i.text

            result = ''.join(map(str, result))

            result = list(result.split(" "))

            result.remove("Pro")
            result.remove("Tips")
            result = ' '.join(result)
            results_list.append(result)
            # print(result)

        if "pro tips" in i.text:
            result = i.text

            result = ''.join(map(str, result))

            result = list(result.split(" "))

            result.remove("pro")
            result.remove("tips")
            result = ' '.join(result)
            results_list.append(result)
            # print(result)


        if "*" in i.text:
            print(i.text)
            result = i.text

            result = ''.join(map(str, result))

            result = result.replace('*','')

            print(result)

            result = list(result.split(" "))

            print(result)

            result = ' '.join(result)
            results_list.append(result)

    
    # if the document actually has pro tips to be written...
    if len(results_list) > 0:


        # print(result) #string
        # print(results_list) #list

        path = '/Users/williameverett/Desktop/Fall-2020/' + class_folder + '/Pro Tips.docx'

        # if pro tip file exists
        if os.path.isfile(path):
            # print("File exists")

            doc_text = docx2txt.process(path)

            document = docx.Document(path)



            # empty list to have string data appended to it
            document_list = []

            # loop through document, add strings to list
            for i in document.paragraphs:
                document_list.append(i.text)


            now = datetime.datetime.now().strftime("%m/%d")


            now = list(now) # make variable a list for formatting purposes

            # remove 0 in month if month < 10, formatting 
            if now[0] == '0': 
                now.remove(now[0])

            now = ''.join(now) #convert date back to a string for formatting


            # print("NOW: ", now)

            # print("DOCUMENT CONTENTS:", document_list)

            # remove date from list
            if now in document_list:
                document_list.remove(now)

            
    


            if results_list[0] not in document_list:

                os.remove(path) #remove pro tip file so that I can create it again and reformat date postings

                #create new protip file
                document = Document()

                header = document.sections[0].header #create header

                header.paragraphs[0].text = class_folder + ' Pro Tips'

                header.paragraphs[0].alignment = 1


                date = document.add_paragraph(now) #add date

                date.alignment = 1 #center dates



                for i in range(len(results_list)):
                    # print(results_list[i])
                    if results_list[i] not in doc_text:
                
                        paragraph = document.add_paragraph(results_list[i], style='ListBullet')
                        paragraph_format = paragraph.paragraph_format

                        paragraph_format.line_spacing = Pt(18)


                # add original elements to pro tip doc
                for i in range(len(document_list)):
                    if any(char.isdigit() for char in document_list[i]) and '/' in document_list[i]:
                        date = document.add_paragraph(document_list[i])
                        date.alignment = 1

                    else:
                        paragraph = document.add_paragraph(document_list[i], style='ListBullet')
                        paragraph_format = paragraph.paragraph_format

                        paragraph_format.line_spacing = Pt(18)



                document.save(path)

        # if pro tip file does not exist
        else:
            document = Document()

            header = document.sections[0].header #create header

            header.paragraphs[0].text = class_folder + ' Pro Tips'

            header.paragraphs[0].alignment = 1



            now = datetime.datetime.now().strftime("%m/%d")


            now = list(now) # make variable a list for formatting purposes

            # remove 0 in month if month < 10, formatting 
            if now[0] == '0': 
                now.remove(now[0])

            now = ''.join(now) #convert date back to a string for formatting


            date = document.add_paragraph(now) #add date

            date.alignment = 1 #center dates


            for i in range(len(results_list)):
                # print(results_list[i])
                paragraph = document.add_paragraph(results_list[i], style='ListBullet')
                paragraph_format = paragraph.paragraph_format

                paragraph_format.line_spacing = Pt(18)



            document.save(path)


if __name__ == "__main__":

    print("\nWelcome to the Pro Tip Generator. Choose your class below or enter 'all' to select all. Enter 'stop' to end program.")

    class_name = input("\n1. OPIM-230\n2. COSC-052\n3. FINC-212\n4. FINC-241\n5. OPIM-258\n\nCourse: ")

    if class_name.upper() == "ALL":

        # class_name = "OPIM-230/Notes.docx"
        # write_to_doc(class_name)

        class_name = "COSC-052/Notes.docx"
        write_to_doc(class_name)

        class_name = "FINC-212/Notes.docx"
        write_to_doc(class_name)

        class_name = "FINC-241/Notes.docx"
        write_to_doc(class_name)

        class_name = "OPIM-258/Notes.docx"
        write_to_doc(class_name)

    else:

        while class_name.upper() != "STOP":

            # if class_name == '1':
            #     class_name = "OPIM-230/Notes.docx"
            #     write_to_doc(class_name)

            if class_name == '2':
                class_name = "COSC-052/Notes.docx"
                write_to_doc(class_name)

            if class_name == '3':
                class_name = "FINC-212/Notes.docx"
                write_to_doc(class_name)

            if class_name == '4':
                class_name = "FINC-241/Notes.docx"
                write_to_doc(class_name)

            if class_name == '5':
                class_name = "OPIM-258/Notes.docx"
                write_to_doc(class_name)

            
            class_name = input("\n1. OPIM-230\n2. COSC-052\n3. FINC-212\n4. FINC-241\n5. OPIM-258\n\nCourse: ")




    print("\nPro Tip generation complete.\n")


